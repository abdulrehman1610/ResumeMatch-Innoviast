"""Resume export module — generates clean, ATS-friendly .docx files with accepted rewrites.

ATS parsers expect:
- Standard section headings (Education, Experience, Skills, etc.)
- Simple single-column layout with no tables, text boxes, or graphics
- Consistent font (Calibri or Arial), 10.5–11pt body, 12–14pt headings
- Bullet points using standard list characters (•), not custom symbols
- Clear hierarchy: Name > Contact > Sections > Bullets
"""

import io
import re
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.schema import RewriteSuggestion


# ---------------------------------------------------------------------------
# ATS-recognized section headings (case-insensitive matching)
# ---------------------------------------------------------------------------
STANDARD_SECTIONS = {
    "summary", "professional summary", "objective", "career objective",
    "profile", "about", "about me",
    "experience", "work experience", "professional experience", "employment",
    "employment history", "work history",
    "education", "academic background", "academic qualifications",
    "skills", "technical skills", "core competencies", "key skills",
    "tools", "technologies", "tech stack",
    "projects", "personal projects", "key projects", "academic projects",
    "certifications", "certificates", "licenses", "credentials",
    "awards", "honors", "achievements", "accomplishments",
    "publications", "research", "papers",
    "volunteer", "volunteering", "community involvement",
    "languages", "interests", "hobbies",
    "references", "additional information",
    "contact", "contact information",
    "internships", "training", "courses", "coursework",
}


def _is_section_heading(line: str) -> bool:
    """
    Determine if a line is a resume section heading using multiple heuristics:
    1. Matches a known standard section name
    2. Is short, starts uppercase, and doesn't end with sentence punctuation
    3. Is ALL CAPS and under 50 chars
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False

    # Remove common decorators: dashes, pipes, colons, underscores, equal signs
    cleaned = re.sub(r'^[\-=_|:•►▪#*]+\s*', '', stripped)
    cleaned = re.sub(r'\s*[\-=_|:]+$', '', cleaned)
    cleaned = cleaned.strip()

    if not cleaned:
        return False

    # Check against known section names
    if cleaned.lower() in STANDARD_SECTIONS:
        return True

    # ALL CAPS short lines are almost always section headers
    if cleaned.upper() == cleaned and len(cleaned) < 50 and cleaned.isalpha():
        return True

    # Title-like: short, no ending period/comma, starts with uppercase letter
    if (
        len(cleaned) < 40
        and not cleaned.endswith(('.', ',', ';'))
        and cleaned[0].isupper()
        and cleaned.lower() in STANDARD_SECTIONS
    ):
        return True

    return False


def _is_bullet_line(line: str) -> bool:
    """Detect if a line is a bullet point."""
    stripped = line.strip()
    # Common bullet prefixes
    return bool(re.match(r'^[\-•●○►▪▸◦∙\*]\s+', stripped)) or bool(re.match(r'^\d+[\.\)]\s+', stripped))


def _clean_bullet_text(line: str) -> str:
    """Strip bullet prefix characters and return clean text."""
    stripped = line.strip()
    cleaned = re.sub(r'^[\-•●○►▪▸◦∙\*]\s+', '', stripped)
    cleaned = re.sub(r'^\d+[\.\)]\s+', '', cleaned)
    return cleaned.strip()


def _standardize_heading(text: str) -> str:
    """Convert heading text to clean title case for ATS consistency."""
    cleaned = re.sub(r'^[\-=_|:•►▪#*]+\s*', '', text.strip())
    cleaned = re.sub(r'\s*[\-=_|:]+$', '', cleaned)
    cleaned = cleaned.strip()

    # Map common variations to standard ATS-recognized headings
    mapping = {
        "work experience": "Professional Experience",
        "employment": "Professional Experience",
        "employment history": "Professional Experience",
        "work history": "Professional Experience",
        "experience": "Professional Experience",
        "professional experience": "Professional Experience",
        "education": "Education",
        "academic background": "Education",
        "academic qualifications": "Education",
        "skills": "Skills",
        "technical skills": "Technical Skills",
        "core competencies": "Core Competencies",
        "key skills": "Key Skills",
        "projects": "Projects",
        "personal projects": "Projects",
        "key projects": "Projects",
        "academic projects": "Projects",
        "certifications": "Certifications",
        "certificates": "Certifications",
        "summary": "Professional Summary",
        "professional summary": "Professional Summary",
        "objective": "Career Objective",
        "career objective": "Career Objective",
        "profile": "Profile",
        "awards": "Awards & Honors",
        "honors": "Awards & Honors",
        "achievements": "Achievements",
        "publications": "Publications",
        "research": "Research",
        "volunteer": "Volunteer Experience",
        "volunteering": "Volunteer Experience",
        "languages": "Languages",
        "interests": "Interests",
        "references": "References",
        "internships": "Internships",
        "training": "Training",
        "courses": "Relevant Coursework",
        "coursework": "Relevant Coursework",
    }

    return mapping.get(cleaned.lower(), cleaned.title())


def _set_paragraph_spacing(paragraph, before_pt: int = 0, after_pt: int = 2, line_spacing_pt: float = 13):
    """Set precise paragraph spacing for tight ATS-friendly layout."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = Pt(line_spacing_pt)


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal line separator (ATS-safe, no graphics)."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before_pt=2, after_pt=2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _detect_contact_block(lines: List[str]) -> Tuple[List[str], List[str]]:
    """
    Separate the top contact block (name, email, phone, links) from the body.
    Heuristic: the first 1–5 lines before the first section heading are contact info.
    """
    contact_lines = []
    body_start = 0

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            if contact_lines:
                body_start = i + 1
                break
            continue

        if _is_section_heading(stripped) and i > 0:
            body_start = i
            break

        # Contact info characteristics: short, contains @, phone patterns, URLs, or is the first non-empty line
        is_contact = (
            i < 6
            and (
                '@' in stripped
                or re.search(r'\b\d{3}[\-.\s]?\d{3,4}[\-.\s]?\d{4}\b', stripped)
                or re.search(r'(linkedin|github|portfolio|http|www\.)', stripped.lower())
                or i == 0  # First line is usually the name
                or len(stripped) < 80
            )
        )

        if is_contact:
            contact_lines.append(stripped)
            body_start = i + 1
        else:
            body_start = i
            break

    return contact_lines, lines[body_start:]


def generate_tailored_resume(
    original_resume_text: str,
    accepted_suggestions: List[RewriteSuggestion],
    candidate_name: str = "Candidate"
) -> bytes:
    """
    Generate a clean ATS-friendly .docx resume with accepted rewrites applied.

    Structure:
    1. Candidate name (large, centered)
    2. Contact info line (centered, smaller)
    3. Horizontal separator
    4. Resume body with properly formatted sections, sub-entries, and bullets
    5. Footer note about applied rewrites
    """
    doc = Document()

    # --- Page setup: standard US Letter with professional margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Apply accepted rewrites to text ---
    modified_text = original_resume_text
    applied_count = 0
    for suggestion in accepted_suggestions:
        if suggestion.original_bullet and suggestion.original_bullet.strip():
            if suggestion.original_bullet in modified_text:
                modified_text = modified_text.replace(
                    suggestion.original_bullet,
                    suggestion.suggested_bullet,
                    1
                )
                applied_count += 1

    # --- Parse lines ---
    all_lines = modified_text.split("\n")
    contact_lines, body_lines = _detect_contact_block(all_lines)

    # --- Render contact header ---
    if contact_lines:
        # First contact line = candidate name
        name_text = contact_lines[0] if contact_lines else candidate_name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(name_para, before_pt=0, after_pt=2, line_spacing_pt=16)
        name_run = name_para.add_run(name_text)
        name_run.font.size = Pt(18)
        name_run.font.name = "Calibri"
        name_run.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Remaining contact lines joined with  |  separator
        if len(contact_lines) > 1:
            contact_text = "  |  ".join(contact_lines[1:])
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(contact_para, before_pt=0, after_pt=4, line_spacing_pt=12)
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = Pt(9.5)
            contact_run.font.name = "Calibri"
            contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    else:
        # Fallback: use provided candidate name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(name_para, before_pt=0, after_pt=4, line_spacing_pt=16)
        name_run = name_para.add_run(candidate_name)
        name_run.font.size = Pt(18)
        name_run.font.name = "Calibri"
        name_run.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    _add_horizontal_rule(doc)

    # --- Render resume body ---
    for line in body_lines:
        stripped = line.strip()

        # Skip empty lines (spacing is handled by paragraph formatting)
        if not stripped:
            continue

        # --- Section heading ---
        if _is_section_heading(stripped):
            heading_text = _standardize_heading(stripped)

            heading_para = doc.add_paragraph()
            _set_paragraph_spacing(heading_para, before_pt=10, after_pt=3, line_spacing_pt=14)
            heading_run = heading_para.add_run(heading_text.upper())
            heading_run.font.size = Pt(11)
            heading_run.font.name = "Calibri"
            heading_run.bold = True
            heading_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            # Add thin line under heading
            pPr = heading_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # --- Bullet point ---
        if _is_bullet_line(stripped):
            bullet_text = _clean_bullet_text(stripped)

            bullet_para = doc.add_paragraph()
            _set_paragraph_spacing(bullet_para, before_pt=0, after_pt=1, line_spacing_pt=12.5)

            # Indent bullets
            pf = bullet_para.paragraph_format
            pf.left_indent = Inches(0.25)

            bullet_run_marker = bullet_para.add_run("•  ")
            bullet_run_marker.font.size = Pt(10.5)
            bullet_run_marker.font.name = "Calibri"
            bullet_run_marker.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            bullet_run_text = bullet_para.add_run(bullet_text)
            bullet_run_text.font.size = Pt(10.5)
            bullet_run_text.font.name = "Calibri"
            bullet_run_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # --- Sub-heading / job title / institution line ---
        # Heuristic: lines that are moderately short, contain dates or pipe/dash separators,
        # or appear right after a section heading — these are entry titles
        is_entry_title = (
            len(stripped) < 120
            and (
                re.search(r'\b(20\d{2}|19\d{2})\b', stripped)  # contains a year
                or '|' in stripped
                or ' – ' in stripped or ' — ' in stripped
                or re.search(r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|present|current)', stripped.lower())
            )
        )

        if is_entry_title:
            entry_para = doc.add_paragraph()
            _set_paragraph_spacing(entry_para, before_pt=6, after_pt=1, line_spacing_pt=13)

            # If there's a date component, try to split title and date
            # Common formats: "Job Title | Company — Date" or "Title, Company (Date)"
            entry_run = entry_para.add_run(stripped)
            entry_run.font.size = Pt(10.5)
            entry_run.font.name = "Calibri"
            entry_run.bold = True
            entry_run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
            continue

        # --- Regular body text ---
        body_para = doc.add_paragraph()
        _set_paragraph_spacing(body_para, before_pt=0, after_pt=2, line_spacing_pt=12.5)
        body_run = body_para.add_run(stripped)
        body_run.font.size = Pt(10.5)
        body_run.font.name = "Calibri"
        body_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Footer: rewrite attribution ---
    if applied_count > 0:
        doc.add_paragraph()  # spacer
        footer_para = doc.add_paragraph()
        _set_paragraph_spacing(footer_para, before_pt=8, after_pt=0)
        footer_run = footer_para.add_run(
            f"— {applied_count} AI-suggested rewrite(s) applied. Review all content before submitting."
        )
        footer_run.font.size = Pt(8)
        footer_run.font.name = "Calibri"
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        footer_run.italic = True

    # --- Export to bytes ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
