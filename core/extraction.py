"""Document extraction module for PDF, DOCX, and raw text files."""

import io
from typing import Union
import pdfplumber
import docx

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5MB cap
MIN_WORD_COUNT = 30  # Minimum word threshold for non-trivial resume text


class ExtractionError(Exception):
    """Raised when text extraction fails or input is corrupt/invalid."""
    pass


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes using pdfplumber."""
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        extracted = "\n".join(text_parts).strip()
        if not extracted:
            raise ExtractionError("PDF file contains no extractable text (it might be scanned/image-only).")
        return extracted
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX file bytes using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table cells if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        extracted = "\n".join(paragraphs).strip()
        if not extracted:
            raise ExtractionError("DOCX file contains no extractable text.")
        return extracted
    except ExtractionError:
        raise
    except Exception as e:
        raise ExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_resume_text(
    file_bytes: bytes,
    file_name: str
) -> str:
    """
    Extract and validate resume text from uploaded bytes based on file extension.
    Raises ExtractionError if file size exceeds limit, format unsupported, or text too short.
    """
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ExtractionError(f"File size exceeds limit of {MAX_FILE_SIZE_BYTES // (1024*1024)}MB.")

    ext = file_name.lower().split(".")[-1]
    
    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        if ext == "doc":
            raise ExtractionError("Legacy .doc format is not supported. Please convert to .docx or .pdf.")
        text = extract_text_from_docx(file_bytes)
    elif ext == "txt":
        text = file_bytes.decode("utf-8", errors="replace").strip()
    else:
        raise ExtractionError(f"Unsupported file extension '.{ext}'. Please upload a PDF, DOCX, or TXT file.")

    validate_extracted_text(text)
    return text


def validate_extracted_text(text: str) -> None:
    """Verify that extracted text contains sufficient words to be a valid resume."""
    word_count = len(text.split())
    if word_count < MIN_WORD_COUNT:
        raise ExtractionError(
            f"Extracted resume text is too short ({word_count} words). "
            f"Please upload a complete resume with at least {MIN_WORD_COUNT} words."
        )
